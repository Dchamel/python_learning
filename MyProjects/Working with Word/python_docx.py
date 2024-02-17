# Material for studying
# https://tokmakov.msk.ru/blog/item/78?ysclid=lr2ryzgb7l672049978

from docx import Document
from docx.shared import Pt, RGBColor

# Part I
document = Document()

style = document.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(14)

# heading1 = document.add_heading('Start of the doc!', 0)
# heading1.alignment = 1

# heading2 = document.add_heading('Paragraph 1', level=1)
# heading2.alignment = 1


# par1 = document.add_paragraph(
#     'Картельные сговоры не допускают ситуации, при которой '
#     'стремящиеся вытеснить традиционное производство, нанотехнологии '
#     'будут обнародованы. Значимость этих проблем настолько очевидна, '
#     'что существующая теория позволяет яркий н выполнить важные задания по '
#     'разработке направлений прогрессивного развития. Однозначно, '
#     'реплицированные с зарубежных источников, современные '
#     'исследования, которые представляют собой яркий пример '
#     'континентально-европейского типа политической культуры, '
#     'будут объединены в целые кластеры себе подобных.')
# par1.alignment = 1
#
# document.save('word.docx')

# Part II
target_part_with_word = 'позволяет яркий н выполн'
target_word = 'яркий'

doc = Document('word.docx')
new_par = ''
for par in doc.paragraphs:
    if target_word in par.text:
        before_str_part = par.text.split(target_part_with_word)[0]
        after_str_part = par.text.split(target_part_with_word)[1]
        before_target = target_part_with_word.split(target_word)[0]
        after_target = target_part_with_word.split(target_word)[1]
        target_word += '<color="red">' + target_word + '</color>'
        new_par = before_str_part + before_target + target_word + after_target + after_str_part
        print(new_par)

        par = par.add_run(new_par)
    #     word = par.add_run(target_word)
    #     red = RGBColor(255, 0, 0)
    #     word.font.color.rgb = red

doc.save('word.docx')
